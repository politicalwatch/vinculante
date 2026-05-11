import csv

import pytest
from docx import Document
from unittest.mock import MagicMock

from vinculante.infrastructure.loaders.docx_loader import DocxLoader, _slugify
from vinculante.presentation.cli.ingest import _PreviewDumpingLoader


def _save(doc, tmp_path, name="test.docx") -> str:
    path = tmp_path / name
    doc.save(str(path))
    return str(path)


# ---------------------------------------------------------------------------
# _slugify helpers
# ---------------------------------------------------------------------------


def test_slugify_basic():
    assert _slugify("EJE 1. EDUCACIÓN") == "eje-1-educacion"


def test_slugify_accents_stripped():
    assert _slugify("1.1 Aportaciones iniciales") == "1-1-aportaciones-iniciales"


def test_slugify_preserves_section_number_boundary():
    assert _slugify("12.1 x") == "12-1-x"


def test_slugify_truncates_at_80():
    long_text = "a" * 100
    assert len(_slugify(long_text)) <= 80


# ---------------------------------------------------------------------------
# Two-column table
# ---------------------------------------------------------------------------


def test_two_col_table_basic(tmp_path):
    doc = Document()
    doc.add_heading("EJE 1. EDUCACIÓN", level=1)
    doc.add_heading("1.1 Aportaciones iniciales", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Título de la propuesta"
    table.cell(0, 1).text = "Desarrollo de la propuesta"
    table.cell(1, 0).text = "Educación emocional"
    table.cell(1, 1).text = "Mejorar la educación emocional en los colegios"
    table.cell(2, 0).text = "Educación sexual"
    table.cell(2, 1).text = "Implementar asignaturas obligatorias"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert len(rows) == 2
    assert rows[0]["topic"] == "EJE 1. EDUCACIÓN"
    assert rows[0]["subtopic"] == "1.1 Aportaciones iniciales"
    assert "Educación emocional" in rows[0]["text"]
    assert "Mejorar la educación emocional" in rows[0]["text"]
    assert rows[0]["reference"] == _slugify("1.1 Aportaciones iniciales Educación emocional")


def test_two_col_header_variant(tmp_path):
    """'Título propuesta' / 'Desarrollo propuesta' header variant is skipped."""
    doc = Document()
    doc.add_heading("EJE 2. EMPLEO", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Título propuesta"
    table.cell(0, 1).text = "Desarrollo propuesta"
    table.cell(1, 0).text = "Empleo digno"
    table.cell(1, 1).text = "Garantizar salarios mínimos dignos"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert len(rows) == 1
    assert rows[0]["topic"] == "EJE 2. EMPLEO"
    assert rows[0]["subtopic"] is None


def test_two_col_text_combines_title_and_body(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Título"
    table.cell(0, 1).text = "Desarrollo"
    table.cell(1, 0).text = "Mi título"
    table.cell(1, 1).text = "Mi cuerpo de propuesta"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert rows[0]["text"] == "Mi título\n\nMi cuerpo de propuesta"


# ---------------------------------------------------------------------------
# Single-column table
# ---------------------------------------------------------------------------


def test_one_col_table(tmp_path):
    doc = Document()
    doc.add_heading("EJE 3. VIVIENDA", level=1)
    doc.add_heading("3.2 Aportaciones Encuentros Territoriales", level=2)
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Desarrollo propuesta"
    table.cell(1, 0).text = "Garantizar el acceso a vivienda asequible para jóvenes"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert len(rows) == 1
    assert rows[0]["text"] == "Garantizar el acceso a vivienda asequible para jóvenes"
    assert rows[0]["topic"] == "EJE 3. VIVIENDA"
    assert rows[0]["subtopic"] == "3.2 Aportaciones Encuentros Territoriales"
    assert rows[0]["reference"].startswith("3-2-aportaciones-encuentros-territoriales")


# ---------------------------------------------------------------------------
# Reference deduplication
# ---------------------------------------------------------------------------


def test_duplicate_titles_get_unique_references(tmp_path):
    doc = Document()
    doc.add_heading("EJE 1. EDUCACIÓN", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Título"
    table.cell(0, 1).text = "Desarrollo"
    table.cell(1, 0).text = "Misma propuesta"
    table.cell(1, 1).text = "Texto A"
    table.cell(2, 0).text = "Misma propuesta"
    table.cell(2, 1).text = "Texto B"

    rows = DocxLoader().load(_save(doc, tmp_path))

    refs = [r["reference"] for r in rows]
    assert len(set(refs)) == 2, "Duplicate titles must produce unique references"
    assert refs[1] == f"{refs[0]}-1"


# ---------------------------------------------------------------------------
# Empty cells / rows skipped
# ---------------------------------------------------------------------------


def test_empty_rows_skipped(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Título"
    table.cell(0, 1).text = "Desarrollo"
    # row 1: both empty
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    # row 2: has content
    table.cell(2, 0).text = "Propuesta válida"
    table.cell(2, 1).text = "Texto de la propuesta"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert len(rows) == 1
    assert rows[0]["text"].startswith("Propuesta válida")


# ---------------------------------------------------------------------------
# Heading context propagates across tables
# ---------------------------------------------------------------------------


def test_heading_context_propagates(tmp_path):
    doc = Document()
    doc.add_heading("EJE 4. SALUD", level=1)
    doc.add_heading("4.1 Propuestas iniciales", level=2)
    t1 = doc.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "Título"
    t1.cell(0, 1).text = "Desarrollo"
    t1.cell(1, 0).text = "Salud mental"
    t1.cell(1, 1).text = "Atención psicológica gratuita"

    doc.add_heading("4.2 Aportaciones Encuentros Territoriales", level=2)
    t2 = doc.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "Título"
    t2.cell(0, 1).text = "Desarrollo"
    t2.cell(1, 0).text = "Salud reproductiva"
    t2.cell(1, 1).text = "Acceso libre a métodos anticonceptivos"

    rows = DocxLoader().load(_save(doc, tmp_path))

    assert len(rows) == 2
    assert rows[0]["subtopic"] == "4.1 Propuestas iniciales"
    assert rows[1]["subtopic"] == "4.2 Aportaciones Encuentros Territoriales"
    assert rows[0]["topic"] == rows[1]["topic"] == "EJE 4. SALUD"


# ---------------------------------------------------------------------------
# _PreviewDumpingLoader
# ---------------------------------------------------------------------------


def test_preview_dumping_loader_writes_csv(tmp_path):
    fake_rows = [
        {"text": "Propuesta A", "reference": "ref-a", "topic": "EJE 1", "subtopic": "1.1"},
        {"text": "Propuesta B", "reference": "ref-b", "topic": "EJE 1", "subtopic": "1.2"},
    ]
    inner = MagicMock()
    inner.load.return_value = fake_rows
    csv_path = tmp_path / "out.preview.csv"

    loader = _PreviewDumpingLoader(inner, csv_path)
    rows = loader.load("dummy.docx")

    assert rows == fake_rows
    assert csv_path.exists()
    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = list(csv.DictReader(f))
    assert len(reader) == 2
    assert reader[0]["reference"] == "ref-a"
    assert reader[1]["topic"] == "EJE 1"
