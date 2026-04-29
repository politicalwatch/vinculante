"""
One-off script: extract citizen proposals and expected matches from raw pilot docx files.
Outputs:
  engine/data/eval/proposals/{domain}.csv
  engine/data/eval/testset.json

Usage:
  uv run --with python-docx scripts/extract_eval.py
"""

import csv
import json
import re
from pathlib import Path

from docx import Document

RAW_DIR = Path(__file__).parent.parent / "data/eval/raw"
PROPOSALS_DIR = Path(__file__).parent.parent / "data/eval/proposals"
TESTSET_PATH = Path(__file__).parent.parent / "data/eval/testset.json"

SOURCES = [
    {
        "docx": RAW_DIR / "PILOTO2_ Proyecto de Ley e insumos_EDUCACION.docx",
        "domain": "educacion",
        "target": "targets/PILOTO2_ Proyecto de Ley e insumos_EDUCACION.pdf",
    },
    {
        "docx": RAW_DIR / "PILOTO2_ Proyecto de Ley e insumos_VIVIENDA.docx",
        "domain": "vivienda",
        "target": "targets/PILOTO2_ Proyecto de Ley e insumos_VIVIENDA.pdf",
    },
]


def extract_proposals(doc: Document) -> list[str]:
    """Return ordered list of citizen proposal texts."""
    proposals = []
    in_proposals = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 1" and "Insumos" in text:
            in_proposals = True
            continue
        if p.style.name == "Heading 1" and in_proposals:
            break
        if "POSIBLES RESULTADOS" in text.upper() and in_proposals:
            break
        if in_proposals:
            proposals.append(text)
    return proposals


def extract_matches(doc: Document, domain: str) -> list[dict]:
    """Return list of {article, proposal_refs} parsed from POSIBLES RESULTADOS."""
    results = []
    in_results = False
    current_article = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if re.match(r"POSIBLES\s+RESULTADOS\s*$", text.upper()):
            in_results = True
            continue

        if not in_results:
            continue

        # Article heading (either Heading 2 or inline at start of paragraph)
        is_article_heading = p.style.name == "Heading 2" and re.match(r"Art[ií]culo\s+\d+", text)
        is_inline_article = p.style.name != "Heading 2" and re.match(r"Art[ií]culo\s+\d+[\.\s]", text)

        if is_article_heading:
            current_article = text
            continue

        if is_inline_article:
            # Article name is the first sentence up to a line break or period after the number
            article_match = re.match(r"(Art[ií]culo\s+\d+\.[^.\n]+)", text)
            current_article = article_match.group(1).strip() if article_match else text.split("\n")[0].strip()

        if current_article and text:
            # Extract insumo numbers from text
            nums = re.findall(r"\b(\d+)\b", text)
            if nums:
                refs = list(dict.fromkeys(f"{domain}-{n}" for n in nums))
                # Merge with existing entry for this article or create new
                existing = next((r for r in results if r["article"] == current_article), None)
                if existing:
                    for ref in refs:
                        if ref not in existing["proposal_refs"]:
                            existing["proposal_refs"].append(ref)
                else:
                    results.append({"article": current_article, "proposal_refs": refs})

            if is_article_heading:
                pass
            elif not is_inline_article:
                # After processing the result text, reset only if next will be a new article
                pass

    return results


def main():
    PROPOSALS_DIR.mkdir(parents=True, exist_ok=True)
    cases = []

    for source in SOURCES:
        domain = source["domain"]
        doc = Document(source["docx"])

        # Extract and write proposals CSV
        proposals = extract_proposals(doc)
        csv_path = PROPOSALS_DIR / f"{domain}.csv"
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["id", "reference", "text"])
            for i, text in enumerate(proposals, start=1):
                writer.writerow([i, f"{domain}-{i}", text])
        print(f"{domain}: {len(proposals)} proposals → {csv_path}")

        # Extract expected matches
        matches = extract_matches(doc, domain)
        cases.append(
            {
                "domain": domain,
                "target": source["target"],
                "proposals": f"proposals/{domain}.csv",
                "expected": matches,
            }
        )
        print(f"{domain}: {len(matches)} article matches extracted")

    testset = {"version": 1, "cases": cases}
    TESTSET_PATH.write_text(json.dumps(testset, ensure_ascii=False, indent=2))
    print(f"testset.json written → {TESTSET_PATH}")


if __name__ == "__main__":
    main()
